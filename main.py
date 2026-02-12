"""
Meeting Intelligence System — FastAPI backend
"""

from pathlib import Path
from uuid import uuid4
import traceback
import importlib

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from starlette.concurrency import run_in_threadpool

from src.recorder import start_recording, stop_recording

from fastapi.responses import FileResponse
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
import os
from datetime import datetime




# ===============================
# Load services ONLY ONCE
# ===============================
def load_services():
    for module_name in ("services", "src.services"):
        try:
            return importlib.import_module(module_name)
        except ModuleNotFoundError:
            continue
    raise RuntimeError("services.py not found")


services = load_services()
process_meeting = getattr(services, "process_meeting")
generate_notes = getattr(services, "generate_notes")
ask_question = getattr(services, "ask_question")


# ===============================
# App setup
# ===============================
app = FastAPI(title="Meeting Intelligence System")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)

ALLOWED_EXTENSIONS = {".mp4", ".mp3", ".wav"}


# ===============================
# Runtime state
# ===============================
app.state.current_meeting_id = None
app.state.meeting_processed = False
stream = None


# ===============================
# Models
# ===============================
class ChatRequest(BaseModel):
    question: str
    meeting_id: str   # 🔥 NEW


# ===============================
# Routes
# ===============================
@app.get("/")
async def root():
    return {"message": "API running"}


# ------------------------------------------------
# Start recording
# ------------------------------------------------
@app.post("/start-recording")
async def start_rec():
    global stream

    try:
        stream = start_recording()
        return {"message": "Recording started"}
    except Exception:
        traceback.print_exc()
        raise HTTPException(500, "Recording failed to start")


# ------------------------------------------------
# Stop recording + process
# ------------------------------------------------
@app.post("/stop-recording")
async def stop_rec():
    global stream

    if stream is None:
        raise HTTPException(400, "Recording not started")

    try:
        meeting_id = uuid4().hex  # 🔥 NEW

        audio_path = stop_recording(stream, "uploads/meeting.wav")

        await run_in_threadpool(process_meeting, str(audio_path), meeting_id)

        app.state.current_meeting_id = meeting_id
        app.state.meeting_processed = True
        stream = None

        return {
            "message": "Recording stopped & meeting processed",
            "meeting_id": meeting_id   # 🔥 send to frontend
        }

    except Exception:
        traceback.print_exc()
        raise HTTPException(500, "Recording processing failed")


# -------------------------------
# Upload + process
# -------------------------------
@app.post("/upload")
async def upload(file: UploadFile = File(...)):

    ext = Path(file.filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(400, "Only mp4/mp3/wav allowed")

    meeting_id = uuid4().hex  # 🔥 NEW

    file_path = UPLOAD_DIR / f"{meeting_id}{ext}"

    try:
        with open(file_path, "wb") as f:
            while chunk := await file.read(1024 * 1024):
                f.write(chunk)

        await run_in_threadpool(process_meeting, str(file_path), meeting_id)

    except Exception:
        traceback.print_exc()
        raise HTTPException(500, "Pipeline crashed. Check terminal.")

    app.state.current_meeting_id = meeting_id
    app.state.meeting_processed = True

    return {
        "message": "meeting processed successfully",
        "meeting_id": meeting_id   # 🔥 send to frontend
    }


# -------------------------------
# Notes
# -------------------------------
@app.post("/notes")
async def notes():

    if not app.state.meeting_processed:
        raise HTTPException(400, "Upload/Record meeting first")

    try:
        # 🔥 pass meeting_id to highlights
        result = await run_in_threadpool(
            generate_notes,
            app.state.current_meeting_id
        )

        return {"notes": result}

    except Exception:
        traceback.print_exc()
        raise HTTPException(500, "Notes generation failed")


# -------------------------------
# Chat
# -------------------------------
@app.post("/chat")
async def chat(payload: ChatRequest):

    if not app.state.meeting_processed:
        raise HTTPException(400, "Upload/Record meeting first")

    try:
        answer = await run_in_threadpool(
            ask_question,
            payload.question,
            payload.meeting_id   # 🔥 NEW
        )

        return {"answer": answer}

    except Exception:
        traceback.print_exc()
        raise HTTPException(500, "Chat failed")




#download notes
# ==============================
# Download Highlights
# ==============================

from fastapi.responses import FileResponse
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
import json


@app.get("/download-notes")
def download_notes(meeting_id: str, format: str = "pdf"):

    # =========================
    # Load Meeting Name
    # =========================

    meeting_name = meeting_id

    mapping_file = "data/meetings.json"

    if os.path.exists(mapping_file):
        with open(mapping_file) as f:
            db = json.load(f)
            meeting_name = db.get(meeting_id, meeting_id)

    # remove unsafe filename characters
    meeting_name = "".join(
        c for c in meeting_name if c.isalnum() or c in (" ", "-", "_")
    ).strip()

    # =========================
    # Load highlights text
    # =========================

    txt_path = f"Notes/highlights_{meeting_id}.txt"

    if not os.path.exists(txt_path):
        return {"error": "Highlights not generated yet."}

    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    # ======================================================
    # ================= PDF DOWNLOAD =======================
    # ======================================================

    if format == "pdf":

        pdf_path = f"Notes/{meeting_name}.pdf"

        doc = SimpleDocTemplate(pdf_path)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph("Meeting Highlights", styles["Heading1"]))
        elements.append(Spacer(1, 20))

        elements.append(Paragraph(f"Meeting: {meeting_name}", styles["Normal"]))
        elements.append(Spacer(1, 20))

        for line in text.split("\n"):
            elements.append(Paragraph(line, styles["BodyText"]))
            elements.append(Spacer(1, 8))

        doc.build(elements)

        return FileResponse(
            pdf_path,
            media_type="application/pdf",
            filename=f"{meeting_name}.pdf"
        )



    # ======================================================
    # ================= TXT DOWNLOAD =======================
    # ======================================================

    elif format == "txt":

        return FileResponse(
            txt_path,
            media_type="text/plain",
            filename=f"{meeting_name}.txt"
        )



    # ======================================================
    # ================= DOCX DOWNLOAD ======================
    # ======================================================

    elif format == "docx":

        docx_path = f"Notes/{meeting_name}.docx"

        document = Document()
        document.add_heading("Meeting Highlights", 0)
        document.add_paragraph(f"Meeting: {meeting_name}")
        document.add_paragraph("")

        for line in text.split("\n"):
            document.add_paragraph(line)

        document.save(docx_path)

        return FileResponse(
            docx_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{meeting_name}.docx"
        )



    # ======================================================
    # ================= INVALID FORMAT =====================
    # ======================================================

    else:
        return {"error": "Invalid format"}



#For set meeting name by user
class MeetingName(BaseModel):
    meeting_id: str
    name: str


@app.post("/set-meeting-name")
def set_meeting_name(data: MeetingName):

    meeting_id = data.meeting_id
    name = data.name

    os.makedirs("data", exist_ok=True)
    file = "data/meetings.json"

    if os.path.exists(file):
        with open(file) as f:
            db = json.load(f)
    else:
        db = {}

    db[meeting_id] = name

    with open(file, "w") as f:
        json.dump(db, f, indent=2)

    return {"status": "saved"}
